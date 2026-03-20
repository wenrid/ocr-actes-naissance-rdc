import streamlit as st
import torch
import re
import unicodedata
from PIL import Image
from transformers import AutoProcessor, AutoModelForImageTextToText
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import sqlite3
import json
from datetime import datetime

st.set_page_config(
    page_title="🇨🇩 Extraction Certificats RDC",
    page_icon="📜",
    layout="wide"
)

DB_PATH = r"C:\Users\rwenc\Desktop\SIM P28\RF\Numerisation_Certificat\actes.db"

def init_db():
    """Crée la table si elle n'existe pas"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS actes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom_enfant TEXT,
            date_extraction TEXT,
            donnees_json TEXT
        )
    ''')
    conn.commit()
    conn.close()

def verifier_doublon(nom: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT id FROM actes WHERE nom_enfant = ?', (nom,))
    existant = c.fetchone()
    conn.close()
    return existant is not None

def sauvegarder_acte(data: dict):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''INSERT INTO actes (nom_enfant, date_extraction, donnees_json)
                 VALUES (?, ?, ?)''',
              (data.get('nom_complet_enfant', ''),
               datetime.now().strftime("%Y-%m-%d %H:%M"),
               json.dumps(data, ensure_ascii=False)))
    conn.commit()
    conn.close()

def remplacer_acte(data: dict):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''UPDATE actes SET date_extraction=?, donnees_json=?
                 WHERE nom_enfant=?''',
              (datetime.now().strftime("%Y-%m-%d %H:%M"),
               json.dumps(data, ensure_ascii=False),
               data.get('nom_complet_enfant', '')))
    conn.commit()
    conn.close()

def rechercher_actes(nom: str) -> list:
    """Recherche par nom de l'enfant"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        SELECT id, nom_enfant, date_extraction, donnees_json
        FROM actes
        WHERE nom_enfant LIKE ?
        ORDER BY date_extraction DESC
    ''', (f'%{nom}%',))
    resultats = c.fetchall()
    conn.close()
    return resultats

def supprimer_acte(acte_id: int):
    """Supprime un acte par son ID"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('DELETE FROM actes WHERE id = ?', (acte_id,))
    conn.commit()
    conn.close()

def stats_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT COUNT(*) FROM actes')
    total = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM actes WHERE date_extraction >= date('now')")
    today = c.fetchone()[0]
    c.execute("SELECT COUNT(*) FROM actes WHERE date_extraction >= date('now','-7 days')")
    week = c.fetchone()[0]
    conn.close()
    return {"total": total, "today": today, "week": week}

# Initialiser la DB au démarrage
init_db()

if 'afficher_tous' not in st.session_state:
    st.session_state['afficher_tous'] = False

# ==================== CONFIGURATION ====================
MODEL_PATH = r"C:\Users\rwenc\Desktop\SIM P28\RF\Numerisation_Certificat\Text_extraction\final5_op\modele_merged_complet"

# ==================== CHARGEMENT DU MODÈLE ====================
@st.cache_resource
def load_model():
    device = "cuda" if torch.cuda.is_available() else "cpu"
    processor = AutoProcessor.from_pretrained(MODEL_PATH, trust_remote_code=True)
    model = AutoModelForImageTextToText.from_pretrained(
        MODEL_PATH,
        dtype=torch.float32,
        device_map=device,
        trust_remote_code=True
    )
    model.eval()
    return model, processor, device

# ==================== POST-CORRECTION ====================
def nettoyer_prediction(texte: str) -> str:
    texte = re.sub(r'^#+\s*', '', texte, flags=re.MULTILINE)
    texte = re.sub(r'\*+', '', texte)
    texte = re.sub(r'^[-•]\s+', '', texte, flags=re.MULTILINE)
    texte = re.sub(r'  +$', '', texte, flags=re.MULTILINE)
    texte = re.sub(r'\n{3,}', '\n\n', texte)
    return texte.strip()

def post_correction(texte: str) -> str:
    CORRECTIONS = {
        # Districts
        "TSITANOV": "TSHANGU", "TCHANGU": "TSHANGU", "TCHANGO": "TSHANGU",
        "TSCHANGU": "TSHANGU", "TCHANGIV": "TSHANGU",
        # Communes
        "NUTILI": "NDJILI", "NDTILI": "NDJILI", "NGTILI": "NDJILI",
        "NDJILLI": "NDJILI", "LINGALA": "LINGWALA", "LINGWALLA": "LINGWALA",
        "NGALLIEMA": "NGALIEMA", "BANDAL": "BANDALUNGWA",
        "KIMBANSEKI": "KIMBANSEKE", "KIMBANTSEKE": "KIMBANSEKE",
        "NGIRIINGIRI": "NGIRI-NGIRI", "NGIRI NGIRI": "NGIRI-NGIRI",
        # Villes
        "SANDATIKA": "GANDAJIKA", "GANDATIKA": "GANDAJIKA",
        "GANDAJICA": "GANDAJIKA", "KISANGAMI": "KISANGANI",
        "KISANGANNI": "KISANGANI", "LUBUMBASKI": "LUBUMBASHI",
        "LUBUMBASI": "LUBUMBASHI", "LUBUMBASHY": "LUBUMBASHI",
        "MBUJIMAYI": "MBUJI-MAYI", "MBUJI MAYI": "MBUJI-MAYI",
        "MWENEDITU": "MWENE-DITU", "MWENE DITU": "MWENE-DITU",
        # Noms
        "KARAO": "KAZADI", "KAZAOI": "KAZADI",
        "NGALELA": "NGALULA", "NGALULLA": "NGALULA",
        "PACALINE": "PASCALINE", "PASCALLINE": "PASCALINE",
        "KASOUGO": "KASONGO", "KASSONGO": "KASONGO",
        "TSHILAMBA": "TSHILOMBO", "TSHILAMBO": "TSHILOMBO",
        "MUKEMDI": "MUKENDI", "MUKENDY": "MUKENDI",
        # Professions
        "Electriclen": "Électricien", "Electicien": "Électricien",
        "Electricien": "Électricien", "Medecin": "Médecin",
        "Infirmierre": "Infirmière", "Pharmaciene": "Pharmacienne",
        "Enseignent": "Enseignant", "Menagere": "Ménagère",
        # Qualités / termes admin
        "PENE": "PÈRE", "PERE": "PÈRE", "MERE": "MÈRE",
        "Officie": "Officier", "ETAT": "État",
        "REPUBLIQUE": "RÉPUBLIQUE", "DEMOCRATIQUE": "DÉMOCRATIQUE",
        "Feminin": "Féminin", "Femmine": "Féminin",
    }
    for err, corr in CORRECTIONS.items():
        texte = texte.replace(err, corr)
    texte = re.sub(r'(\d{3})1(\d{3})', r'\1/\2', texte)
    texte = re.sub(r'\b([IVX]+)(\d{4})\b', r'\1/\2', texte)
    return texte

# ==================== EXTRACTION OCR ====================
def extraire_texte(model, processor, device, image: Image.Image) -> str:
    image = image.convert("RGB").resize((1240, 1754), Image.LANCZOS)

    conversation = [{
        "role": "user",
        "content": [
            {"type": "image", "image": image},
            {"type": "text",  "text": "Transcris exactement le texte de ce document sans markdown."}
        ]
    }]

    inputs = processor.apply_chat_template(
        conversation, add_generation_prompt=True,
        tokenize=True, return_dict=True, return_tensors="pt"
    )

    # ← NE PAS caster les inputs, laisser le modèle gérer
    inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        output_ids = model.generate(
            **inputs,
            max_new_tokens=1024,
            do_sample=False,
            temperature=None,
            top_p=None,
        )

    print(f"Output shape: {output_ids.shape}")
    generated_ids = output_ids[0, inputs["input_ids"].shape[1]:]
    print(f"Generated tokens: {len(generated_ids)}")

    texte = processor.decode(generated_ids, skip_special_tokens=True)
    print(f"Texte: '{texte[:100]}'")

    print(f"BRUT AVANT NETTOYAGE: '{texte[:200]}'")

    texte = nettoyer_prediction(texte)
    texte = post_correction(texte)
    return texte


# ==================== PARSER LE TEXTE BRUT ====================
def parser_texte(texte: str) -> dict:
    """Extrait les champs structurés depuis le texte OCR brut"""
    data = {
        "province": "", "ville": "", "district": "", "territoire": "",
        "chefferie": "", "bureau_principal": "", "bureau_secondaire": "",
        "numero_acte": "", "volume": "", "folio": "",
        "annee_declaration": "", "jour_declaration": "", "mois_declaration": "",
        "heure_declaration": "", "minutes_declaration": "",
        "nom_officier": "", "commune_officier": "",
        "nom_declarant": "", "qualite_declarant": "",
        "lieu_naissance_declarant": "", "date_naissance_declarant": "",
        "profession_declarant": "", "residence_declarant": "",
        "jour_naissance_enfant": "", "mois_naissance_enfant": "",
        "annee_naissance_enfant": "", "lieu_naissance_enfant": "",
        "sexe_enfant": "", "nom_complet_enfant": "",
        "nom_pere": "", "lieu_naissance_pere": "", "date_naissance_pere": "",
        "nationalite_pere": "", "profession_pere": "", "residence_pere": "",
        "nom_mere": "", "lieu_naissance_mere": "", "date_naissance_mere": "",
        "nationalite_mere": "", "profession_mere": "", "residence_parents": "",
        "langue_traduction": "Français"
    }

    def extraire(pattern, texte, groupe=1):
        m = re.search(pattern, texte, re.IGNORECASE)
        return m.group(groupe).strip() if m else ""

    def nettoyer_valeur(valeur: str) -> str:
        """Supprime les préfixes parasites capturés par l'OCR"""
        valeur = re.sub(r"L'[EÉ]TAT CIVIL DE\s*:\s*", "", valeur, flags=re.IGNORECASE)
        valeur = re.sub(r"\bDE\s*:\s*", "", valeur, flags=re.IGNORECASE)
        valeur = re.sub(r"\s{2,}", " ", valeur)  # espaces multiples
        return valeur.strip()

    def nettoyer_nom(nom: str) -> str:
        """Supprime les espaces parasites dans les noms"""
        # Supprime espaces entre lettres majuscules collées (ex: TSHI KALA → TSHIKALA)
        nom = re.sub(r'([A-Z]{2,})\s([A-Z]{2,})', lambda m: m.group(0), nom)
        nom = re.sub(r'\s{2,}', ' ', nom)
        return nom.strip()

    # ==================== EN-TÊTE ====================
    data["province"]  = extraire(r"Province de\s*:?\s*([A-ZÉÈÊ\-]+)", texte)
    data["ville"]     = extraire(r"Ville de\s*:?\s*([A-ZÉÈÊ\-]+)", texte)
    data["district"]  = extraire(r"District de\s*:?\s*([A-ZÉÈÊ\-]+)", texte)
    data["territoire"]= extraire(r"Territoire\s*/.*?de\s*:?\s*([A-ZÉÈÊ\-]+)", texte)
    data["chefferie"] = extraire(r"Chefferie.*?de\s*:?\s*([A-ZÉÈÊ\-]+)", texte)

    data["bureau_principal"]  = nettoyer_valeur(extraire(r"Bureau Principal.*?de\s*:?\s*(.+)", texte))
    data["bureau_secondaire"] = nettoyer_valeur(extraire(r"Bureau secondaire.*?de\s*:?\s*(.+)", texte))

    # ==================== ACTE / VOLUME / FOLIO ====================
    m_acte = re.search(r"Acte\s*n[°o]?\s*:?\s*([\d/]+)", texte, re.IGNORECASE)
    if m_acte:
        data["numero_acte"] = m_acte.group(1).strip()

    # Volume : cherche I/2026 ou II/2026 etc.
    m_vol = re.search(r"Volume\s*:?\s*([IVX]+/\d{4})", texte, re.IGNORECASE)
    if m_vol:
        data["volume"] = m_vol.group(1).strip()
    else:
        # Tentative alternative : Volume suivi de texte
        m_vol2 = re.search(r"Volume\s*:?\s*(\S+)", texte, re.IGNORECASE)
        if m_vol2:
            data["volume"] = m_vol2.group(1).strip()

    m_folio = re.search(r"Folio\s*n[°o]?\s*:?\s*(\d+)", texte, re.IGNORECASE)
    if m_folio:
        data["folio"] = m_folio.group(1).strip()

    # ==================== DÉCLARATION ====================
    m_ann = re.search(r"deux mille\s+([A-ZÉÈÊÀ\-]+)", texte, re.IGNORECASE)
    if m_ann:
        data["annee_declaration"] = m_ann.group(1).strip()

    m_jour = re.search(r"deux mille\s+\S+\s+le\s+([A-ZÉÈÊÀ\-]+(?:\s+[A-ZÉÈÊÀ\-]+)?)\s+jour", texte, re.IGNORECASE)
    if m_jour:
        data["jour_declaration"] = m_jour.group(1).strip()

    m_mois = re.search(r"jour du mois\s+de\s+([A-ZÉÈÊÀ]+)", texte, re.IGNORECASE)
    if m_mois:
        data["mois_declaration"] = m_mois.group(1).strip()

    m_heure = re.search(r"[àa]\s+(\d{1,2})\s+heures?\s+(\d{0,2})", texte, re.IGNORECASE)
    if m_heure:
        data["heure_declaration"]   = m_heure.group(1).strip()
        data["minutes_declaration"] = m_heure.group(2).strip()

    # ==================== OFFICIER ====================
    # "Par devant nous * LUMUMBA OKITU SERGE"
    m_off = re.search(r"Par devant nous\s*\*?\s*([A-ZÉÈÊ\s]+?)(?:\n|Officier|$)", texte, re.IGNORECASE)
    if m_off:
        data["nom_officier"] = m_off.group(1).strip()

    # Commune officier — après "Officier de l'État civil de"
    m_comm = re.search(r"Officier de l'[EÉ]tat civil de\s+([A-ZÉÈÊ\s]+?)(?:\n|A comparu|$)", texte, re.IGNORECASE)
    if m_comm:
        val = m_comm.group(1).strip()
        # Nettoie "LA COMMUNE DE" si présent
        val = re.sub(r"LA COMMUNE DE\s*", "", val, flags=re.IGNORECASE)
        data["commune_officier"] = val.strip()

    # ==================== DÉCLARANT ====================
    m_dec = re.search(r"A comparu\s+(.+?)\s+en qualit[eé] de\s+(\S+)", texte, re.IGNORECASE)
    if m_dec:
        data["nom_declarant"]     = nettoyer_nom(m_dec.group(1).strip())
        data["qualite_declarant"] = m_dec.group(2).strip()

    m_nais_dec = re.search(r"N[eé][eé]?\s*[àa]\*?\s+(.+?)\s+le\s+([\d/]+)", texte, re.IGNORECASE)
    if m_nais_dec:
        data["lieu_naissance_declarant"] = m_nais_dec.group(1).strip()
        data["date_naissance_declarant"] = m_nais_dec.group(2).strip()

    # Profession : limiter jusqu'au mot suivant (RÉSIDENT, NÉ, ET DE...)
    m_prof_dec = re.search(
        r"Profession\s+(.+?)(?:\s+R[EÉ]SIDENT|\s+NÉ|\s+ET DE|\s+NATIONALITÉ|\n|$)",
        texte, re.IGNORECASE
    )
    if m_prof_dec:
        data["profession_declarant"] = m_prof_dec.group(1).strip()

    m_res_dec = re.search(r"R[eé]sident[eé]?\s+[àaÀ]\s+(.+?)(?:\n|$)", texte, re.IGNORECASE)
    if m_res_dec:
        data["residence_declarant"] = m_res_dec.group(1).strip()

    # ==================== ENFANT ====================
    m_enf_date = re.search(
        r"Le\s+([A-ZÉÈÊÀ\-]+(?:\s+[A-ZÉÈÊÀ\-]+)?)\s+jour du mois de\s+([A-ZÉÈÊÀ]+)\s+de l'ann[eé]e\s+(\d{4})",
        texte, re.IGNORECASE
    )
    if m_enf_date:
        data["jour_naissance_enfant"]  = m_enf_date.group(1).strip()
        data["mois_naissance_enfant"]  = m_enf_date.group(2).strip()
        data["annee_naissance_enfant"] = m_enf_date.group(3).strip()

    m_enf_lieu = re.search(r"est n[eé]\s+[àaÀ]\s+([A-ZÉÈÊ\-]+)\s+un enfant de sexe\s+(\S+)", texte, re.IGNORECASE)
    if m_enf_lieu:
        data["lieu_naissance_enfant"] = m_enf_lieu.group(1).strip()
        data["sexe_enfant"]           = m_enf_lieu.group(2).strip()

    m_nom_enf = re.search(r"nomm[eé]\s+([A-ZÉÈÊ\s]+?)(?:\n|fils|$)", texte, re.IGNORECASE)
    if m_nom_enf:
        data["nom_complet_enfant"] = nettoyer_nom(m_nom_enf.group(1).strip())

    # ==================== PÈRE ====================
    # "fils (fille) de* NOM né à LIEU"
    m_pere = re.search(
        r"fils \(fille\) de\*?\s+(.+?)\s+n[eé]\s+[àaÀ]\s+([A-ZÉÈÊ\-]+)",
        texte, re.IGNORECASE
    )
    if m_pere:
        data["nom_pere"]            = nettoyer_nom(m_pere.group(1).strip())
        data["lieu_naissance_pere"] = m_pere.group(2).strip()

    # Date, nationalité, profession père
    m_pere2 = re.search(
        r"(?:fils.*?n[eé].*?[àa]\s+\S+\s+)?[Ll][Ee]\s+([\d/]+)\s+[Nn]ationalit[eé]\s+(\S+)\s+[Pp]rofession\s+(.+?)\s+[Rr][eé]sident",
        texte, re.IGNORECASE
    )
    if m_pere2:
        data["date_naissance_pere"] = m_pere2.group(1).strip()
        data["nationalite_pere"]    = m_pere2.group(2).strip()
        data["profession_pere"]     = m_pere2.group(3).strip()

    m_res_pere = re.search(
        r"[Rr][eé]sident\s+[àaÀ]\s+(.+?)\s+et de",
        texte, re.IGNORECASE
    )
    if m_res_pere:
        data["residence_pere"] = m_res_pere.group(1).strip()

    # ==================== MÈRE ====================
    # "et de NOM né(e) à LIEU"
    m_mere = re.search(
        r"et de\s+([A-ZÉÈÊ\s]+?)\s+n[eé][eé]?\s*\([eé]\)\s+[àaÀ]\s+([A-ZÉÈÊ\-]+)",
        texte, re.IGNORECASE
    )
    if m_mere:
        data["nom_mere"]            = nettoyer_nom(m_mere.group(1).strip())
        data["lieu_naissance_mere"] = m_mere.group(2).strip()

    # Date, nationalité, profession mère (deuxième occurrence)
    matches_natio = re.findall(
        r"[Ll][Ee]\s+([\d/]+)\s+[Nn]ationalit[eé]\s+(\S+)\s+[Pp]rofession\s+(.+?)\s+[Rr][eé]sident",
        texte, re.IGNORECASE
    )
    if len(matches_natio) >= 2:
        data["date_naissance_mere"] = matches_natio[1][0].strip()
        data["nationalite_mere"]    = matches_natio[1][1].strip()
        data["profession_mere"]     = matches_natio[1][2].strip()

    m_res_par = re.search(
        r"[Rr][eé]sidents\s+[àaÀ]\s+(.+?)\s+[Cc]onjoints",
        texte, re.IGNORECASE
    )
    if m_res_par:
        data["residence_parents"] = m_res_par.group(1).strip()

    # ==================== LANGUE ====================
    if re.search(r"Lingala", texte, re.IGNORECASE):
        data["langue_traduction"] = "Lingala"
    elif re.search(r"Swahili", texte, re.IGNORECASE):
        data["langue_traduction"] = "Swahili"

    # ==================== POST-NETTOYAGE ====================
    # Ajouter MONT-AMBA dans post_correction
    data["district"] = data["district"].replace("MONT-ANBA", "MONT-AMBA")

    return data

# ==================== GÉNÉRATION WORD ====================
def generer_acte_naissance_rdc(data: dict) -> BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

    # En-tête
    ht = doc.add_table(rows=1, cols=2)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    cl = ht.cell(0, 0)
    cl.width = Cm(12)
    p = cl.paragraphs[0]
    r = p.add_run("RÉPUBLIQUE DÉMOCRATIQUE DU CONGO")
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.bold = True
    cr = ht.cell(0, 1)
    cr.width = Cm(5)
    p2 = cr.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Volet")
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.font.italic = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    def add_field(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        r1 = p.add_run(label)
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(10)
        r2 = p.add_run(str(value).upper() if value else "")
        r2.font.name = 'Segoe Script'; r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0, 0, 139)

    def add_mixed(parts):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        for txt, script in parts:
            r = p.add_run(str(txt) if txt else "")
            r.font.name = 'Segoe Script' if script else 'Times New Roman'
            r.font.size  = Pt(10)
            if script: r.font.color.rgb = RGBColor(0, 0, 139)

    # Champs en-tête
    add_field("Province de : ", data['province'])
    add_field("Ville de : ", data['ville'])
    add_field("District de : ", data['district'])
    add_field("Territoire /Secteur ou Cité de : ", data['territoire'])
    add_field("Chefferie/secteur ou Cité de : ", data['chefferie'])
    add_field("Bureau Principal de l'État civil de : ", data['bureau_principal'])
    add_field("Bureau secondaire de l'État civil de : ", data['bureau_secondaire'])

    # Acte/Volume/Folio sur une ligne
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    for lbl, key in [("Acte n° : ", "numero_acte"), (" Volume : ", "volume"), (" Folio n° : ", "folio")]:
        r1 = p.add_run(lbl); r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
        r2 = p.add_run(data.get(key, ''))
        r2.font.name = 'Segoe Script'; r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0, 0, 139)

    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)

    # Titre
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run("ACTE DE NAISSANCE")
    rt.font.name = 'Times New Roman'; rt.font.size = Pt(16)
    rt.font.bold = True; rt.font.underline = True

    doc.add_paragraph()

    # Date déclaration
    add_mixed([
        ("L'an deux mille ", False),
        (data['annee_declaration'].lower(), True),
        (" le ", False),
        (data['jour_declaration'].lower(), True),
        (" jour du mois", False),
    ])
    add_mixed([
        ("de ", False),
        (data['mois_declaration'].lower(), True),
        (" à ", False),
        (data['heure_declaration'], True),
        (" heures ", False),
        (data['minutes_declaration'], True),
    ])

    # Officier
    add_field("Par devant nous * ", data['nom_officier'])
    add_field("Officier de l'État civil de ", data['commune_officier'])

    # Déclarant
    add_mixed([
        ("A comparu ", False),
        (data['nom_declarant'].upper(), True),
        (" en qualité de ", False),
        (data['qualite_declarant'].upper(), True),
    ])
    add_mixed([
        ("Né(e) à* ", False),
        (data['lieu_naissance_declarant'].upper(), True),
        (" le ", False),
        (data['date_naissance_declarant'], True),
    ])
    add_field("Profession ", data['profession_declarant'])
    add_field("Résident à ", data['residence_declarant'])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Le quel (quelle) nous a déclaré ce qui suit :").font.name = 'Times New Roman'

    # Enfant
    add_mixed([
        ("Le ", False),
        (data['jour_naissance_enfant'].lower(), True),
        (" jour du mois de ", False),
        (data['mois_naissance_enfant'].lower(), True),
        (" de l'année ", False),
        (data['annee_naissance_enfant'], True),
    ])
    add_mixed([
        ("est né à ", False),
        (data['lieu_naissance_enfant'].upper(), True),
        (" un enfant de sexe ", False),
        (data['sexe_enfant'], True),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("nommé "); r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
    r2 = p.add_run(data['nom_complet_enfant'].upper())
    r2.font.name = 'Segoe Script'; r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0, 0, 139); r2.font.bold = True

    # Père
    add_mixed([
        ("fils (fille) de* ", False),
        (data['nom_pere'].upper(), True),
        (" né à ", False),
        (data['lieu_naissance_pere'].upper(), True),
    ])
    add_mixed([
        ("le ", False),
        (data['date_naissance_pere'], True),
        (" nationalité ", False),
        (data['nationalite_pere'], True),
        (" profession ", False),
        (data['profession_pere'], True),
    ])
    add_mixed([
        ("résident à ", False),
        (data['residence_pere'], True),
        (" et de ", False),
    ])

    # Mère
    add_mixed([
        (data['nom_mere'].upper(), True),
        (" né(e) à ", False),
        (data['lieu_naissance_mere'].upper(), True),
    ])
    add_mixed([
        ("Le ", False),
        (data['date_naissance_mere'], True),
        (" nationalité ", False),
        (data['nationalite_mere'], True),
        (" profession ", False),
        (data['profession_mere'], True),
    ])
    add_mixed([
        ("résidents à ", False),
        (data['residence_parents'], True),
        (" conjoints.", False),
    ])

    # Lecture
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Lecture de l'acte a été faite ou connaissance de l'acte a été donnée ou").font.name = 'Times New Roman'

    add_mixed([
        ("traduction de l'acte a été faite en ", False),
        (data['langue_traduction'], True),
        (" langue que nous", False),
    ])

    p = doc.add_paragraph()
    p.add_run("connaissons ou par interprète ayant prêté serment").font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run("En foi de quoi, avons dressé le présent acte.").font.name = 'Times New Roman'

    # Signatures
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    st_table = doc.add_table(rows=1, cols=2)
    st_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cl = st_table.cell(0, 0)
    r = cl.paragraphs[0].add_run("Le déclarant")
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True
    cr = st_table.cell(0, 1)
    p_sig = cr.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p_sig.add_run("L'Officier de l'État civil")
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(*)Préciser le nom et qualité\n(*)Offet les itérations initaires")
    r.font.name = 'Times New Roman'; r.font.size = Pt(8); r.font.italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ==================== INTERFACE ====================
col_flag, col_titre = st.columns([1, 6])
with col_flag:
    st.image(r"C:\Users\rwenc\Desktop\SIM P28\RF\Numerisation_Certificat\Text_extraction\rdc_flag.jpg", width=300)
with col_titre:
    st.title("🇨🇩 Extraction des Certificats de Naissance")
    st.markdown("### République Démocratique du Congo")
st.markdown("---")

with st.spinner("⏳ Chargement du modèle LightOnOCR..."):
    model, processor, device = load_model()
st.success(f"✅ Modèle chargé sur {device}")

stats = stats_db()
m1, m2, m3, m4 = st.columns(4)

with m1:
    st.markdown(f"""
    <div style="background:white;border-radius:12px;padding:18px 20px;
                border-left:4px solid #007FFF;box-shadow:0 2px 8px rgba(0,0,0,0.08)">
        <div style="font-size:2rem;font-weight:700;color:#007FFF">{stats['total']}</div>
        <div style="font-size:0.75rem;color:#5A6A8A;font-weight:600;
                    text-transform:uppercase;letter-spacing:0.05em">🗂️ Actes en base</div>
    </div>""", unsafe_allow_html=True)

with m2:
    st.markdown(f"""
    <div style="background:white;border-radius:12px;padding:18px 20px;
                border-left:4px solid #22C55E;box-shadow:0 2px 8px rgba(0,0,0,0.08)">
        <div style="font-size:2rem;font-weight:700;color:#22C55E">{stats['today']}</div>
        <div style="font-size:0.75rem;color:#5A6A8A;font-weight:600;
                    text-transform:uppercase;letter-spacing:0.05em">📅 Extraits aujourd'hui</div>
    </div>""", unsafe_allow_html=True)

with m3:
    st.markdown(f"""
    <div style="background:white;border-radius:12px;padding:18px 20px;
                border-left:4px solid #F7D918;box-shadow:0 2px 8px rgba(0,0,0,0.08)">
        <div style="font-size:2rem;font-weight:700;color:#F7D918">{stats['week']}</div>
        <div style="font-size:0.75rem;color:#5A6A8A;font-weight:600;
                    text-transform:uppercase;letter-spacing:0.05em">📈 Cette semaine</div>
    </div>""", unsafe_allow_html=True)

with m4:
    nom_recherche = st.text_input("Rechercher", placeholder="🔎 Ex: TSHIMANGA ou SAMUEL", 
                                   label_visibility="collapsed")
st.markdown("---")


if st.button("📋 Voir tous les certificats dans la base de données", use_container_width=True):
    st.session_state['afficher_tous'] = True

if st.session_state.get('afficher_tous'):
    resultats = rechercher_actes("")
    if resultats:
        st.success(f"✅ {len(resultats)} acte(s) en base de données")
        for acte_id, nom, date_ext, donnees_json in resultats:
            with st.expander(f"📄 {nom} — extrait le {date_ext}"):
                data = json.loads(donnees_json)
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**👶 Enfant**")
                    st.write(f"Nom : {data.get('nom_complet_enfant', '')}")
                    st.write(f"Né(e) le : {data.get('jour_naissance_enfant', '')} {data.get('mois_naissance_enfant', '')} {data.get('annee_naissance_enfant', '')}")
                    st.write(f"Lieu : {data.get('lieu_naissance_enfant', '')}")
                with col2:
                    st.markdown("**👨👩 Parents**")
                    st.write(f"Père : {data.get('nom_pere', '')}")
                    st.write(f"Mère : {data.get('nom_mere', '')}")

                    
                col_dl, col_mod, col_del = st.columns(3)
                with col_dl:
                    buf = generer_acte_naissance_rdc(data)
                    st.download_button("📥 Télécharger Word", data=buf,
                        file_name=f"acte_{nom.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_tous_{acte_id}",
                        use_container_width=True)

                with col_mod:
                    if st.button("✏️ Modifier", key=f"mod_tous_{acte_id}",
                                  use_container_width=True, type="secondary"):
                        st.session_state[f'modifier_tous_{acte_id}'] = True

                with col_del:
                    if st.button("🗑️ Supprimer", key=f"del_tous_{acte_id}",
                                  use_container_width=True, type="primary"):
                        supprimer_acte(acte_id)
                        st.session_state['afficher_tous'] = False
                        st.rerun()

                if st.session_state.get(f'modifier_tous_{acte_id}'):
                    st.markdown("---")
                    st.markdown("**✏️ Modifier les données**")
                    c1, c2 = st.columns(2)
                    with c1:
                        new_nom  = st.text_input("👶 Nom enfant",     data.get('nom_complet_enfant',''), key=f"new_nom_tous_{acte_id}")
                        new_pere = st.text_input("👨 Nom père",       data.get('nom_pere',''),           key=f"new_pere_tous_{acte_id}")
                        new_lieu = st.text_input("📍 Lieu naissance", data.get('lieu_naissance_enfant',''), key=f"new_lieu_tous_{acte_id}")
                    with c2:
                        new_mere  = st.text_input("👩 Nom mère",  data.get('nom_mere',''),               key=f"new_mere_tous_{acte_id}")
                        new_sexe  = st.text_input("⚥ Sexe",       data.get('sexe_enfant',''),            key=f"new_sexe_tous_{acte_id}")
                        new_annee = st.text_input("📅 Année",      data.get('annee_naissance_enfant',''), key=f"new_annee_tous_{acte_id}")

                    if st.button("💾 Sauvegarder", key=f"save_tous_{acte_id}", type="primary"):
                        data.update({
                            'nom_complet_enfant': new_nom,
                            'nom_pere': new_pere,
                            'nom_mere': new_mere,
                            'sexe_enfant': new_sexe,
                            'lieu_naissance_enfant': new_lieu,
                            'annee_naissance_enfant': new_annee,
                        })
                        conn = sqlite3.connect(DB_PATH)
                        c = conn.cursor()
                        c.execute('UPDATE actes SET nom_enfant=?, donnees_json=? WHERE id=?',
                                  (new_nom, json.dumps(data, ensure_ascii=False), acte_id))
                        conn.commit(); conn.close()
                        st.session_state[f'modifier_tous_{acte_id}'] = False
                        st.success("✅ Modifications sauvegardées !")
                        st.rerun()

    else:
        st.warning("Aucun acte en base de données.")

col1, col2 = st.columns(2)

with col1:
    st.subheader("📤 Télécharger le certificat")
    uploaded = st.file_uploader(
        "Sélectionner un acte de naissance",
        type=['png', 'jpg', 'jpeg'],
        label_visibility="collapsed"
    )

    if uploaded:
        image = Image.open(uploaded)
        st.image(image, caption="Certificat uploadé", use_container_width=True)

        if st.button("🔍 EXTRAIRE LE TEXTE", type="primary", use_container_width=True):
            progress = st.progress(0, text="⏳ Démarrage...")
            progress.progress(10, text="🖼️ Prétraitement de l'image...")
            texte_brut = extraire_texte(model, processor, device, image)
            progress.progress(60, text="🤖 Reconnaissance OCR en cours...")
            parsed = parser_texte(texte_brut)
            progress.progress(90, text="🔎 Vérification en base de données...")
            st.session_state['texte_brut'] = texte_brut
            st.session_state['parsed']     = parsed

            nom = parsed.get('nom_complet_enfant', '')
            if verifier_doublon(nom):
                st.session_state['doublon_detecte'] = True
                st.session_state['doublon_nom']     = nom
            else:
                sauvegarder_acte(parsed)
                progress.progress(100, text="✅ Terminé !")
                st.success("✅ Extraction terminée !")
                st.info("💾 Acte sauvegardé dans la base de données !")

            with st.expander("📄 Texte OCR brut"):
                st.text(texte_brut)

        if st.session_state.get('doublon_detecte'):
            nom = st.session_state.get('doublon_nom', '')
            st.warning(f"⚠️ **{nom}** existe déjà dans la base de données. Voulez-vous le remplacer ?")
            col_oui, col_non = st.columns(2)
            with col_oui:
                if st.button("✅ Oui, remplacer", use_container_width=True, type="primary"):
                    remplacer_acte(st.session_state['parsed'])
                    st.session_state['doublon_detecte'] = False
                    st.success("✅ Acte remplacé avec succès !")
                    st.rerun()
            with col_non:
                if st.button("❌ Non, annuler", use_container_width=True):
                    st.session_state['doublon_detecte'] = False
                    st.info("ℹ️ Enregistrement annulé.")
                    st.rerun()    

with col2:
    st.subheader("📝 Données extraites")

    if 'parsed' in st.session_state:
        parsed = st.session_state['parsed']

        with st.expander("✏️ Modifier les données", expanded=True):
            edited = {}
            st.markdown("**👶 Enfant**")
            edited['nom_complet_enfant']    = st.text_input("Nom complet enfant",    parsed.get('nom_complet_enfant', ''))
            edited['sexe_enfant']           = st.text_input("Sexe",                  parsed.get('sexe_enfant', ''))
            edited['lieu_naissance_enfant'] = st.text_input("Lieu naissance enfant", parsed.get('lieu_naissance_enfant', ''))
            edited['jour_naissance_enfant'] = st.text_input("Jour naissance",        parsed.get('jour_naissance_enfant', ''))
            edited['mois_naissance_enfant'] = st.text_input("Mois naissance",        parsed.get('mois_naissance_enfant', ''))
            edited['annee_naissance_enfant']= st.text_input("Année naissance",       parsed.get('annee_naissance_enfant', ''))

            st.markdown("**👨 Père**")
            edited['nom_pere']        = st.text_input("Nom père",        parsed.get('nom_pere', ''))
            edited['profession_pere'] = st.text_input("Profession père", parsed.get('profession_pere', ''))

            st.markdown("**👩 Mère**")
            edited['nom_mere']        = st.text_input("Nom mère",        parsed.get('nom_mere', ''))
            edited['profession_mere'] = st.text_input("Profession mère", parsed.get('profession_mere', ''))

            st.markdown("**🏛️ Administratif**")
            edited['nom_officier']   = st.text_input("Officier",         parsed.get('nom_officier', ''))
            edited['nom_declarant']  = st.text_input("Déclarant",        parsed.get('nom_declarant', ''))
            edited['numero_acte']    = st.text_input("N° Acte",          parsed.get('numero_acte', ''))

            final_data = {**parsed, **edited}
            st.session_state['final_data'] = final_data

        st.markdown("---")
        st.subheader("📥 Export Word")

        if st.button("📄 GÉNÉRER LE DOCUMENT WORD", type="primary", use_container_width=True):
            final = st.session_state.get('final_data', parsed)
            buf   = generer_acte_naissance_rdc(final)
            nom_fichier = f"acte_{final.get('nom_complet_enfant','extrait').replace(' ', '_')}.docx"
            st.download_button(
                "📥 TÉLÉCHARGER L'ACTE DE NAISSANCE",
                data=buf,
                file_name=nom_fichier,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
st.markdown("---")
st.subheader("🔎 Certificat(s) trouvé(s)")


if nom_recherche:
    resultats = rechercher_actes(nom_recherche)

    if resultats:
        st.success(f"✅ {len(resultats)} acte(s) trouvé(s)")

        for acte_id, nom, date_ext, donnees_json in resultats:
            is_open = st.session_state.get(f'modifier_{acte_id}', False)
            with st.expander(f"📄 {nom} — extrait le {date_ext}", expanded=is_open):
                data = json.loads(donnees_json)

                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**👶 Enfant**")
                    st.write(f"Nom : {data.get('nom_complet_enfant', '')}")
                    st.write(f"Né(e) le : {data.get('jour_naissance_enfant', '')} {data.get('mois_naissance_enfant', '')} {data.get('annee_naissance_enfant', '')}")
                    st.write(f"Lieu : {data.get('lieu_naissance_enfant', '')}")
                    st.write(f"Sexe : {data.get('sexe_enfant', '')}")

                with col2:
                    st.markdown("**👨👩 Parents**")
                    st.write(f"Père : {data.get('nom_pere', '')}")
                    st.write(f"Mère : {data.get('nom_mere', '')}")
                    st.write(f"Résidence : {data.get('residence_parents', '')}")

                col_dl, col_mod, col_del = st.columns(3)
                with col_dl:
                    buf = generer_acte_naissance_rdc(data)
                    st.download_button("📥 Télécharger Word", data=buf,
                        file_name=f"acte_{nom.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{acte_id}",
                        use_container_width=True)

                with col_mod:
                    if st.button("✏️ Modifier", key=f"mod_{acte_id}",
                                  use_container_width=True, type="secondary"):
                        st.session_state[f'modifier_{acte_id}'] = True

                with col_del:
                    if st.button("🗑️ Supprimer", key=f"del_{acte_id}",
                                  use_container_width=True, type="primary"):
                        supprimer_acte(acte_id)
                        st.rerun()

                if st.session_state.get(f'modifier_{acte_id}'):
                    st.markdown("---")
                    st.markdown("**✏️ Modifier les données**")
                    c1, c2 = st.columns(2)
                    with c1:
                        new_nom  = st.text_input("👶 Nom enfant",     data.get('nom_complet_enfant',''), key=f"new_nom_{acte_id}")
                        new_pere = st.text_input("👨 Nom père",       data.get('nom_pere',''),           key=f"new_pere_{acte_id}")
                        new_lieu = st.text_input("📍 Lieu naissance", data.get('lieu_naissance_enfant',''), key=f"new_lieu_{acte_id}")
                    with c2:
                        new_mere  = st.text_input("👩 Nom mère",  data.get('nom_mere',''),               key=f"new_mere_{acte_id}")
                        new_sexe  = st.text_input("⚥ Sexe",       data.get('sexe_enfant',''),            key=f"new_sexe_{acte_id}")
                        new_annee = st.text_input("📅 Année",      data.get('annee_naissance_enfant',''), key=f"new_annee_{acte_id}")

                    if st.button("💾 Sauvegarder", key=f"save_{acte_id}", type="primary"):
                        data.update({
                            'nom_complet_enfant': new_nom,
                            'nom_pere': new_pere,
                            'nom_mere': new_mere,
                            'sexe_enfant': new_sexe,
                            'lieu_naissance_enfant': new_lieu,
                            'annee_naissance_enfant': new_annee,
                        })
                        conn = sqlite3.connect(DB_PATH)
                        c = conn.cursor()
                        c.execute('UPDATE actes SET nom_enfant=?, donnees_json=? WHERE id=?',
                                  (new_nom, json.dumps(data, ensure_ascii=False), acte_id))
                        conn.commit(); conn.close()
                        st.session_state[f'modifier_{acte_id}'] = False
                        st.success("✅ Modifications sauvegardées !")
                        st.rerun()

    else:
        st.warning(f"Aucun acte trouvé pour '{nom_recherche}'")

st.markdown("---")
st.markdown("""
<div style="text-align:center;color:gray;font-size:0.85rem">
    <em> Numérisation des Actes de Naissance RDC</em>
</div>
""", unsafe_allow_html=True)